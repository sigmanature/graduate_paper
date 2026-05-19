#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
pdf_translate_to_docx.py

端到端：解析 PDF（含图片）→（调用你自己的大模型 API 翻译）→ 按 Word 模板样式生成译文 docx。

✅ 特点
- 尽量保留论文“行文结构顺序”：标题/作者/摘要/关键词/分节/小节/图表题注/参考文献
- 支持两栏论文的阅读顺序（左栏→右栏，并在跨栏块处做分段）
- 支持提取 PDF 图片并插入 Word（近似按出现顺序）
- 翻译步骤可接入 OpenAI Responses API，也可替换成任何“兼容接口”
- 生成中间 JSON（elements + translations）支持断点续跑

依赖：
  pip install pymupdf python-docx pillow openai

用法示例：
  export OPENAI_API_KEY="你的key"
  python pdf_translate_to_docx.py \
      --pdf "paper.pdf" \
      --template "模板.docx" \
      --out "译文.docx" \
      --provider openai \
      --model "gpt-5.2" \
      --cache "cache_translations.json"

如果你要用别的模型/供应商：
- 直接改 Translator.translate_marked() 内部的 HTTP 调用即可（保持输入输出格式一致）
"""

from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import fitz  # PyMuPDF
from PIL import Image  # noqa: F401 (docx add_picture 需要 pillow 支持某些格式)
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.shared import Pt

# ---------------------------
# 1) 结构化提取：PDF -> elements
# ---------------------------

def _sha1(text: str) -> str:
    return hashlib.sha1(text.encode("utf-8")).hexdigest()

def cleanup_text(s: str) -> str:
    """合并断字、压缩空白、去掉多余换行（块内）。"""
    if not s:
        return ""
    # archi-\ntectures -> architectures
    s = re.sub(r"(\w)-\s*\n\s*(\w)", r"\1\2", s)
    # 块内多行变空格
    s = re.sub(r"\s*\n\s*", " ", s)
    # 多空格压缩
    s = re.sub(r"[ \t]{2,}", " ", s)
    return s.strip()

def normalize_spaced_caps_heading(s: str) -> str:
    """
    PyMuPDF 有时会把全大写单词拆开：'I. I NTRODUCTION'。
    这个函数把类似 'I NTRODUCTION' -> 'INTRODUCTION'。
    """
    s = s.strip()
    # 反复把 “单个大写字母 + 空格 + 2个以上大写” 拼起来
    for _ in range(5):
        ns = re.sub(r"\b([A-Z])\s+([A-Z]{2,})\b", r"\1\2", s)
        if ns == s:
            break
        s = ns
    return s

def looks_like_section_heading(text: str) -> Optional[Tuple[int, str]]:
    """
    返回 (level, heading_text)
    level: 1 = 大节(I., II., ...), 2 = 小节(A., B., ...)
    """
    t = normalize_spaced_caps_heading(text)
    # 大节：罗马数字
    m = re.match(r"^([IVX]+)\.\s+(.+)$", t)
    if m:
        return 1, f"{m.group(1)}. {m.group(2).strip()}"
    # 小节：A. / B. / C.
    m = re.match(r"^([A-Z])\.\s+(.+)$", t)
    if m:
        return 2, f"{m.group(1)}. {m.group(2).strip()}"
    return None

def looks_like_caption(text: str) -> bool:
    t = text.strip()
    return bool(re.match(r"^(Fig\.|Figure|Table)\s*\d+[:\.]?\s+", t))

def looks_like_references_heading(text: str) -> bool:
    return text.strip().upper().startswith("REFERENCES")

def avg_font_size_of_block(block: dict) -> float:
    sizes = []
    for line in block.get("lines", []):
        for span in line.get("spans", []):
            if "size" in span:
                sizes.append(span["size"])
    return sum(sizes) / len(sizes) if sizes else 0.0

def block_text(block: dict) -> str:
    """把一个 text block 的 lines/spans 合并成字符串（保留行序，但最终会 cleanup）。"""
    if block.get("type") != 0:
        return ""
    parts = []
    for line in block.get("lines", []):
        line_text = "".join(span.get("text", "") for span in line.get("spans", []))
        parts.append(line_text)
    return "\n".join(parts)

def is_full_width(bbox: Tuple[float, float, float, float], page_width: float, ratio: float = 0.82) -> bool:
    x0, _, x1, _ = bbox
    return (x1 - x0) >= page_width * ratio

def order_blocks_reading(blocks: List[dict], page_width: float) -> List[dict]:
    """
    近似阅读顺序（支持两栏）：
    - 把跨栏(full-width)块按 y 排序，当作分隔点；
    - 每个分段内：先左栏自上而下，再右栏自上而下；
    """
    # 过滤掉空 bbox
    blocks = [b for b in blocks if b.get("bbox")]
    # 先按 y,x 粗排，便于取统计
    blocks_sorted = sorted(blocks, key=lambda b: (b["bbox"][1], b["bbox"][0]))

    text_blocks = [b for b in blocks_sorted if b.get("type") == 0]
    if len(text_blocks) < 8:
        return blocks_sorted  # 页内容太少，就不做两栏处理

    # 判断两栏：看 x0 是否明显分成左右两团
    x0s = [b["bbox"][0] for b in text_blocks]
    mid = page_width / 2.0
    left = [b for b in blocks_sorted if b["bbox"][0] < mid and not is_full_width(b["bbox"], page_width)]
    right = [b for b in blocks_sorted if b["bbox"][0] >= mid and not is_full_width(b["bbox"], page_width)]
    two_col = (len(left) >= 5 and len(right) >= 5)

    if not two_col:
        return blocks_sorted

    full = [b for b in blocks_sorted if is_full_width(b["bbox"], page_width)]
    full = sorted(full, key=lambda b: (b["bbox"][1], b["bbox"][0]))

    remaining = [b for b in blocks_sorted if b not in full]
    out: List[dict] = []

    def emit_segment(y0: float, y1: float):
        seg = [b for b in remaining if (b["bbox"][1] >= y0 and b["bbox"][1] < y1)]
        seg_left = sorted([b for b in seg if b["bbox"][0] < mid], key=lambda b: (b["bbox"][1], b["bbox"][0]))
        seg_right = sorted([b for b in seg if b["bbox"][0] >= mid], key=lambda b: (b["bbox"][1], b["bbox"][0]))
        out.extend(seg_left)
        out.extend(seg_right)

    cur_y = -1e9
    for fb in full:
        fy0, fy1 = fb["bbox"][1], fb["bbox"][3]
        emit_segment(cur_y, fy0)
        out.append(fb)
        cur_y = fy1

    emit_segment(cur_y, 1e12)

    # 去重保持顺序
    seen = set()
    uniq = []
    for b in out:
        k = (b.get("number"), b.get("bbox")[0], b.get("bbox")[1], b.get("type"))
        if k in seen:
            continue
        seen.add(k)
        uniq.append(b)
    return uniq

@dataclass
class Element:
    type: str  # title/author/affiliation/abstract/keywords/heading/subheading/paragraph/caption/image/reference
    text: Optional[str] = None
    page: Optional[int] = None
    bbox: Optional[Tuple[float, float, float, float]] = None
    level: Optional[int] = None
    image_path: Optional[str] = None
    translate: bool = True

def extract_pdf_elements(pdf_path: Path, image_dir: Path) -> List[Element]:
    image_dir.mkdir(parents=True, exist_ok=True)
    doc = fitz.open(str(pdf_path))

    elements: List[Element] = []
    in_references = False
    seen_first_page_header_done = False

    for page_index in range(len(doc)):
        page = doc[page_index]
        page_dict = page.get_text("dict")
        blocks = page_dict.get("blocks", [])
        ordered = order_blocks_reading(blocks, page.rect.width)

        for b in ordered:
            bbox = tuple(b.get("bbox", (0, 0, 0, 0)))
            if b.get("type") == 1:
                # 图片块
                img_bytes = b.get("image")
                ext = (b.get("ext") or "png").lower()
                if not img_bytes:
                    continue
                img_name = f"p{page_index+1:02d}_{len([e for e in elements if e.type=='image'])+1:03d}.{ext}"
                img_path = image_dir / img_name
                img_path.write_bytes(img_bytes)
                elements.append(Element(type="image", page=page_index + 1, bbox=bbox, image_path=str(img_path), translate=False))
                continue

            if b.get("type") != 0:
                continue

            raw = block_text(b).strip()
            if not raw:
                continue
            avg_size = avg_font_size_of_block(b)
            txt = cleanup_text(raw)

            # 参考文献段落：不翻译（默认）
            if looks_like_references_heading(txt):
                in_references = True
                elements.append(Element(type="heading", text="REFERENCES", page=page_index + 1, bbox=bbox, level=1, translate=False))
                continue

            # 论文开头（标题/作者/单位/邮箱/摘要/关键词）识别：主要针对第一页
            if page_index == 0 and not seen_first_page_header_done and not in_references:
                # 标题一般字体更大、且在 abstract 前
                # 简化：在遇到 Abstract 之前，按行位置依次分类（更稳的做法是看 bbox.y）
                if txt.startswith("Abstract"):
                    seen_first_page_header_done = True
                    # 把 Abstract 这一块单独处理（Abstract + 余下段落）
                    # 例：'Abstract —In this work ...'
                    elements.append(Element(type="abstract", text=txt, page=1, bbox=bbox, translate=True))
                    continue

                # 第一页开头的若干块，大概率是 title/author/affiliation/email
                if avg_size >= 14.0 and len(txt) <= 120:
                    elements.append(Element(type="title", text=txt, page=1, bbox=bbox, translate=True))
                else:
                    # 经验规则：包含 @ 或 .edu/.in 等通常是 email
                    if "@" in txt:
                        elements.append(Element(type="author", text=txt, page=1, bbox=bbox, translate=False))
                    else:
                        elements.append(Element(type="affiliation", text=txt, page=1, bbox=bbox, translate=True))
                continue

            # Keywords 识别（一般在 Abstract 后）
            if txt.startswith("Keywords"):
                elements.append(Element(type="keywords", text=txt, page=page_index + 1, bbox=bbox, translate=True))
                continue

            # 图表题注
            if looks_like_caption(txt):
                elements.append(Element(type="caption", text=txt, page=page_index + 1, bbox=bbox, translate=True))
                continue

            # 分节/小节
            sec = looks_like_section_heading(txt)
            if sec:
                level, heading_txt = sec
                elements.append(Element(
                    type="heading" if level == 1 else "subheading",
                    text=heading_txt,
                    page=page_index + 1,
                    bbox=bbox,
                    level=level,
                    translate=True
                ))
                continue

            # 参考文献内容
            if in_references:
                elements.append(Element(type="reference", text=txt, page=page_index + 1, bbox=bbox, translate=False))
                continue

            # 普通段落
            elements.append(Element(type="paragraph", text=txt, page=page_index + 1, bbox=bbox, translate=True))

    doc.close()
    # 进一步：把 title/author/affiliation 分组合并（可选）。这里保持最保守：按抽取顺序输出。
    return elements

# ---------------------------
# 2) 翻译：elements -> translated elements
# ---------------------------

TRANSLATE_INSTRUCTIONS = """你是一名严谨的英文→中文学术译者。请把输入翻译为中文，并遵守：
1) 保留原有编号/层级：例如 “I. INTRODUCTION” 译为 “I. 引言”，A./B. 小节同理。
2) 保留缩写与专有名词（TLB, MMU, x86-64, MARSSx86, QEMU 等），必要时首次出现可加中文解释括号。
3) 保留引用标号与符号（如 [12]、% 、GB、ns、µs、(a)(b) 等）。
4) 不要改写或扩写含义；优先忠实、清晰、技术风格。
5) 只翻译每个标记块正文，严格保留标记行：<<<P0>>> 这类行必须原样输出。
输出格式：按输入顺序，逐块输出：标记行 + 翻译后的内容。不要添加额外解释。
"""

class Translator:
    """
    你可以把这里替换成任何大模型调用：
    - OpenAI Responses API（默认实现）
    - 兼容 OpenAI 的第三方（改 base_url）
    - 你自己的 HTTP 服务
    """
    def __init__(self, provider: str, model: str, api_key: Optional[str], base_url: Optional[str] = None, temperature: float = 0.2):
        self.provider = provider
        self.model = model
        self.api_key = api_key
        self.base_url = base_url
        self.temperature = temperature
        self._client = None

        if provider == "openai":
            try:
                from openai import OpenAI  # type: ignore
            except Exception as e:
                raise RuntimeError("未安装 openai 包。请 pip install openai") from e
            kwargs = {}
            if api_key:
                kwargs["api_key"] = api_key
            if base_url:
                kwargs["base_url"] = base_url
            self._client = OpenAI(**kwargs)

    def translate_marked(self, marked_text: str) -> str:
        """
        输入：包含多个 <<<P{i}>>> 块的字符串
        输出：同样包含 <<<P{i}>>> 的翻译字符串
        """
        if self.provider == "openai":
            # OpenAI Python SDK 示例见官方文档：client.responses.create(...)。
            # https://developers.openai.com/api/reference/python  citeturn0search6
            resp = self._client.responses.create(
                model=self.model,
                instructions=TRANSLATE_INSTRUCTIONS,
                input=marked_text,
                temperature=self.temperature,
            )
            return resp.output_text

        raise NotImplementedError(f"provider={self.provider} 未实现。")

def parse_marked_output(output: str) -> Dict[int, str]:
    """
    把模型输出解析成 {i: translated_text}
    允许块之间空行，但要求标记行存在。
    """
    pattern = re.compile(r"<<</?P(\d+)>>>|<<<P(\d+)>>>")  # 容错（不会用到）
    # 主要用：<<<P123>>>
    blocks = re.split(r"(<<<P\d+>>>)", output)
    out: Dict[int, str] = {}
    cur_idx = None
    buf: List[str] = []
    for part in blocks:
        m = re.match(r"^<<<P(\d+)>>>$", part.strip())
        if m:
            # flush
            if cur_idx is not None:
                out[cur_idx] = cleanup_text("\n".join(buf))
            cur_idx = int(m.group(1))
            buf = []
        else:
            if cur_idx is not None:
                buf.append(part)
    if cur_idx is not None:
        out[cur_idx] = cleanup_text("\n".join(buf))
    return out

def batch_translate_texts(translator: Translator, texts: List[str], max_chars: int = 9000, sleep_s: float = 0.0) -> List[str]:
    """
    把 texts 分批拼成 marked_text 调用 translate_marked，再拆回列表。
    max_chars 控制单次请求大小（按字符粗略限制，避免超长）。
    """
    results: List[str] = [""] * len(texts)
    i = 0
    while i < len(texts):
        batch = []
        batch_ids = []
        char_count = 0
        while i < len(texts):
            t = texts[i]
            # 空段落直接跳过
            if not t.strip():
                results[i] = ""
                i += 1
                continue
            chunk = f"<<<P{len(batch)}>>>\n{t.strip()}\n"
            if batch and char_count + len(chunk) > max_chars:
                break
            batch.append(chunk)
            batch_ids.append(i)
            char_count += len(chunk)
            i += 1

        marked_text = "\n".join(batch)
        translated = translator.translate_marked(marked_text)
        mapping = parse_marked_output(translated)

        # 写回 results
        for local_j, global_i in enumerate(batch_ids):
            results[global_i] = mapping.get(local_j, "").strip()

        if sleep_s > 0:
            time.sleep(sleep_s)

    return results

def translate_elements(elements: List[Element], translator: Optional[Translator], cache_path: Path, keep_references_english: bool = True) -> List[Element]:
    """
    把需要翻译的 element.text 翻译成中文，写回 element.text。
    使用 cache 避免重复计费。
    """
    cache: Dict[str, str] = {}
    if cache_path.exists():
        cache = json.loads(cache_path.read_text(encoding="utf-8"))

    # 收集待翻译文本
    idxs = []
    texts = []
    for idx, e in enumerate(elements):
        if not e.text or not e.translate:
            continue
        if keep_references_english and e.type in ("reference",):
            continue
        key = _sha1(e.type + "\n" + e.text)
        if key in cache:
            continue
        idxs.append(idx)
        texts.append(e.text)

    if texts:
        if translator is None:
            raise RuntimeError("存在待翻译文本，但未配置 translator。")
        translated_list = batch_translate_texts(translator, texts)
        for idx, zh in zip(idxs, translated_list):
            key = _sha1(elements[idx].type + "\n" + (elements[idx].text or ""))
            cache[key] = zh

        cache_path.write_text(json.dumps(cache, ensure_ascii=False, indent=2), encoding="utf-8")

    # 写回 elements（用 cache）
    out = []
    for e in elements:
        if e.text and e.translate and not (keep_references_english and e.type == "reference"):
            key = _sha1(e.type + "\n" + e.text)
            zh = cache.get(key, e.text)
            out.append(Element(**{**e.__dict__, "text": zh}))
        else:
            out.append(e)
    return out

# ---------------------------
# 3) Word 写入：模板 -> 输出 docx
# ---------------------------

def remove_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None  # type: ignore

def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: Optional[str] = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        try:
            new_para.style = style
        except Exception:
            pass
    if text:
        new_para.add_run(text)
    return new_para

def set_paragraph_basic_format(p: Paragraph, align: Optional[int] = None, bold: Optional[bool] = None, font_size_pt: Optional[float] = None) -> None:
    if align is not None:
        p.alignment = align
    if p.runs:
        r = p.runs[0]
    else:
        r = p.add_run("")
    if bold is not None:
        r.bold = bold
    if font_size_pt is not None:
        r.font.size = Pt(font_size_pt)

def insert_image_after(paragraph: Paragraph, image_path: str, max_width_emu: int) -> Paragraph:
    p = insert_paragraph_after(paragraph, text="", style="Normal")
    run = p.add_run()
    run.add_picture(image_path, width=max_width_emu)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def build_docx_from_template(template_path: Path, out_path: Path, elements: List[Element], source_line: str = "") -> None:
    doc = Document(str(template_path))

    # 找锚点：外文资料翻译 与 原文出处
    start_idx = None
    end_idx = None
    for i, p in enumerate(doc.paragraphs):
        if start_idx is None and "外文资料翻译" in (p.text or ""):
            start_idx = i
        if "原文出处" in (p.text or ""):
            end_idx = i
            break

    if start_idx is None or end_idx is None or end_idx <= start_idx:
        raise RuntimeError("未在模板中找到 '外文资料翻译' 或 '原文出处' 段落，请检查模板内容。")

    # 删除 start_idx 和 end_idx 之间的内容（不删锚点自身）
    to_delete = doc.paragraphs[start_idx + 1:end_idx]
    for p in list(to_delete):
        remove_paragraph(p)

    # 计算图片最大宽度（EMU）
    section = doc.sections[0]
    max_width = section.page_width - section.left_margin - section.right_margin

    # 插入译文内容
    cursor = doc.paragraphs[start_idx]
    for e in elements:
        if e.type == "image" and e.image_path:
            cursor = insert_image_after(cursor, e.image_path, max_width)
            continue

        if not e.text:
            continue

        if e.type == "title":
            cursor = insert_paragraph_after(cursor, e.text, style="Heading 1")
            set_paragraph_basic_format(cursor, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, font_size_pt=16)
        elif e.type in ("author", "affiliation"):
            cursor = insert_paragraph_after(cursor, e.text, style="Normal")
            set_paragraph_basic_format(cursor, align=WD_ALIGN_PARAGRAPH.CENTER)
        elif e.type == "abstract":
            # 分两段：标题 + 内容（abstract 块里可能带 'Abstract —'）
            cursor = insert_paragraph_after(cursor, "摘要", style="Heading 2")
            abs_text = re.sub(r"^Abstract\s*[—\-:]?\s*", "", e.text.strip(), flags=re.IGNORECASE)
            cursor = insert_paragraph_after(cursor, abs_text, style="Normal")
        elif e.type == "keywords":
            cursor = insert_paragraph_after(cursor, "关键词", style="Heading 2")
            kw_text = re.sub(r"^Keywords\s*[—\-:]?\s*", "", e.text.strip(), flags=re.IGNORECASE)
            cursor = insert_paragraph_after(cursor, kw_text, style="Normal")
        elif e.type == "heading":
            cursor = insert_paragraph_after(cursor, e.text, style="Heading 1")
        elif e.type == "subheading":
            cursor = insert_paragraph_after(cursor, e.text, style="Heading 2")
        elif e.type == "caption":
            cursor = insert_paragraph_after(cursor, e.text, style="Normal")
            # 题注常用居中
            set_paragraph_basic_format(cursor, align=WD_ALIGN_PARAGRAPH.CENTER)
        elif e.type == "reference":
            # 若你想把 references 放到模板中的“参考文献”处，可在这里加一个“参考文献”标题。
            cursor = insert_paragraph_after(cursor, e.text, style="Normal")
        else:
            cursor = insert_paragraph_after(cursor, e.text, style="Normal")

    # 写原文出处
    end_para = doc.paragraphs[end_idx]  # '原文出处：'
    # 模板通常下一段是出处内容
    if end_idx + 1 < len(doc.paragraphs) and source_line:
        doc.paragraphs[end_idx + 1].text = source_line

    doc.save(str(out_path))

# ---------------------------
# main
# ---------------------------

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--pdf", required=True, help="输入 PDF 路径")
    ap.add_argument("--template", required=True, help="参考模板 docx 路径")
    ap.add_argument("--out", required=True, help="输出 docx 路径")
    ap.add_argument("--workdir", default="work_pdf2docx", help="中间文件目录（images/elements.json 等）")
    ap.add_argument("--provider", default="openai", help="翻译供应商：openai / 其他(需自己实现)")
    ap.add_argument("--model", default="gpt-5.2", help="翻译模型名")
    ap.add_argument("--api-key", default=None, help="API key（也可用环境变量 OPENAI_API_KEY）")
    ap.add_argument("--base-url", default=None, help="兼容 OpenAI 的第三方 base_url（可选）")
    ap.add_argument("--cache", default="cache_translations.json", help="翻译缓存 json")
    ap.add_argument("--no-translate", action="store_true", help="只抽取结构+生成 docx（不翻译）")
    ap.add_argument("--keep-references-english", action="store_true", help="参考文献保持英文（默认）")
    ap.add_argument("--sleep", type=float, default=0.0, help="每次翻译请求后 sleep 秒数（防限流）")
    ap.add_argument("--source-line", default="", help="写入模板中“原文出处”下一段的文本，例如 'arXiv:2002.01073'")
    args = ap.parse_args()

    pdf_path = Path(args.pdf)
    template_path = Path(args.template)
    out_path = Path(args.out)
    workdir = Path(args.workdir)
    image_dir = workdir / "images"
    workdir.mkdir(parents=True, exist_ok=True)

    elements = extract_pdf_elements(pdf_path, image_dir)
    (workdir / "elements.json").write_text(
        json.dumps([e.__dict__ for e in elements], ensure_ascii=False, indent=2),
        encoding="utf-8"
    )

    translator = None
    if not args.no_translate:
        api_key = args.api_key or os.environ.get("OPENAI_API_KEY")
        translator = Translator(
            provider=args.provider,
            model=args.model,
            api_key=api_key,
            base_url=args.base_url,
        )

    translated = elements
    if not args.no_translate:
        translated = translate_elements(
            elements,
            translator=translator,
            cache_path=Path(args.cache),
            keep_references_english=True,
        )
        (workdir / "elements_translated.json").write_text(
            json.dumps([e.__dict__ for e in translated], ensure_ascii=False, indent=2),
            encoding="utf-8"
        )

    build_docx_from_template(
        template_path=template_path,
        out_path=out_path,
        elements=translated,
        source_line=args.source_line,
    )

    print(f"OK: wrote {out_path}")
    print(f"debug: {workdir}/elements.json, {workdir}/elements_translated.json, {image_dir}/")

if __name__ == "__main__":
    main()
