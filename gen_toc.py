from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER, WD_LINE_SPACING
from docx.oxml.ns import qn

def create_perfect_toc():
    # --- 1. 初始化文档与页面 ---
    doc = Document()
    section = doc.sections[0]

    # ★ 修改默认“正文/Normal”样式为：宋体 小四
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(12)  # 小四
    # 设置 A4 纸张与边距 (标准公文设定)
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    margin_x = Cm(3.54) # 左右边距 2.54cm (注：此处保留了你代码中的3.54，如需标准2.54可自行修改)
    section.left_margin = margin_x
    section.right_margin = margin_x

    # 计算制表位的位置（页宽 - 左右边距）
    max_width = section.page_width - section.left_margin - section.right_margin

    # --- 2. 核心功能函数：添加一行目录 ---
    def add_row(text, page_num,
                level=0,                 # 缩进级别 (0, 1, 2...)
                font_name='黑体',        # 标题字体
                font_size=Pt(14),        # 标题字号
                page_font='宋体',        # 页码字体
                is_bold=False):          # 是否加粗

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT # 强制左对齐

        # [A] 精确计算缩进
        # Level 0 (章): 0 缩进
        # Level 1 (节): 缩进 2 字符 (约 24pt)
        # Level 2 (小节): 缩进 4 字符 (约 48pt)
        indent_step = 24
        p.paragraph_format.left_indent = Pt(level * indent_step)
        # ★ 行距相关：1.5倍行距 + 段前段后 0
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

        # [B] 设置制表位 (关键步骤)
        # 在行尾添加一个“右对齐”的桩子，并指定前导符为点 (DOTS)
        # 注意：python-docx 的 WD_TAB_LEADER.DOTS 对应 Word 中的点
        tabs = p.paragraph_format.tab_stops
        tabs.add_tab_stop(max_width, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.MIDDLE_DOT)

        # [C] 写入：标题文字
        run_txt = p.add_run(text)
        run_txt.font.name = font_name
        run_txt._element.rPr.rFonts.set(qn('w:eastAsia'), font_name) # 强制中文字体
        run_txt.font.size = font_size
        run_txt.bold = is_bold

        # [D] 写入：前导点 (......) -> 核心难点解决
        # 这里的关键是：虽然 Tab 只是一个制表符，但我们可以给它单独设置字体！
        # 设为“宋体”可以让点看起来位置居中；设为“小四(12pt)”符合您的要求。
        run_tab = p.add_run("\t")
        run_tab.font.name = '宋体'
        run_tab._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run_tab.font.size = Pt(12) # 宋体 小四
        run_tab.font.color.rgb = RGBColor(0, 0, 0)

        # [E] 写入：页码
        run_page = p.add_run(page_num)
        run_page.font.name = page_font # 灵活控制页码字体
        run_page._element.rPr.rFonts.set(qn('w:eastAsia'), page_font)
        run_page.font.size = Pt(12) # 页码字号通常为小四
        run_page.bold = is_bold

    # --- 3. 生成文档内容 ---

    # >>> 标题部分：目录 (宋体 三号 居中 不加粗) <<<
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = p_title.add_run("目  录")
    run_t.font.name = '宋体'
    run_t._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run_t.font.size = Pt(16) # 三号
    run_t.bold = False       # 明确不加粗

    # 空一行
    doc.add_paragraph()

    # >>> 数据配置区：在这里修改或添加章节 <<<
    # 常用字号常量
    S_4   = Pt(14) # 四号
    S_X4  = Pt(12) # 小四

    # 数据列表：每一行代表目录的一条
    # 格式：(标题, 页码, 层级, 标题字体, 标题字号, 页码字体)
    # 规则：
    # Level 0 (章): 黑体, 四号(S_4), 页码字体 Times New Roman (除了摘要用宋体)
    # Level 1 (节): 宋体, 小四(S_X4), 页码字体 Times New Roman
    # Level 2 (小节): 宋体, 小四(S_X4), 页码字体 Times New Roman

    toc_data = [
        # 摘要：黑体 四号，页码 I (宋体)
        {"txt": "摘要", "pg": "I", "lvl": 0, "fn": "黑体", "fs": S_4, "pg_fn": "宋体"},

        # ABSTRACT：Times New Roman 四号，页码 II (宋体)
        {"txt": "ABSTRACT", "pg": "II", "lvl": 0, "fn": "Times New Roman", "fs": S_4, "pg_fn": "宋体"},

        # 第一章
        {"txt": "第一章 研究背景", "pg": "1", "lvl": 0, "fn": "黑体", "fs": S_4, "pg_fn": "宋体"},

        {"txt": "1.1 内存管理机制的演进", "pg": "1", "lvl": 1, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
            {"txt": "1.1.1 页表与 TLB 的工作原理", "pg": "1", "lvl": 2, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
            {"txt": "1.1.2 4KB 页面在大内存系统中的扩展性问题", "pg": "1", "lvl": 2, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
        {"txt": "1.2 Linux 文件系统 I/O 栈概述", "pg": "3", "lvl": 1, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
            {"txt": "1.2.1 Page Cache 与 VFS 层交互", "pg": "3", "lvl": 2, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
            {"txt": "1.2.2 传统 Buffer Head 映射机制及其局限性", "pg": "5", "lvl": 2, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
        {"txt": "1.3 F2FS 文件系统架构", "pg": "6", "lvl": 1, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
            {"txt": "1.3.1 日志结构文件系统 (LFS) 与闪存感知设计", "pg": "7", "lvl": 2, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
            {"txt": "1.3.2 数据布局、Segment 结构与 Checkpoint 机制", "pg": "8", "lvl": 2, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},

        # 第二章
        {"txt": "第二章 引入 Folio 和 iomap 的意义", "pg": "10", "lvl": 0, "fn": "黑体", "fs": S_4, "pg_fn": "宋体"},

        {"txt": "2.1 Struct Folio：新一代内存管理单元", "pg": "10", "lvl": 1, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
            {"txt": "2.1.1 Linux 传统大页机制回顾：hugetlbfs、透明大页（THP）与复合页", "pg": "10", "lvl": 2, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
            {"txt": "2.1.2 Folio 在内核中的演进现状与趋势", "pg": "12", "lvl": 2, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
        {"txt": "2.2 iomap 框架：现代化的 I/O 路径", "pg": "14", "lvl": 1, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
            {"txt": "2.2.1 iomap 数据结构与接口定义", "pg": "14", "lvl": 2, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
            {"txt": "2.2.2 迭代器模式与区间映射优势", "pg": "17", "lvl": 2, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},

        # 第三章
        {"txt": "第三章 F2FS 支持 Large Folios 的核心难点", "pg": "18", "lvl": 0, "fn": "黑体", "fs": S_4, "pg_fn": "宋体"},

        {"txt": "3.1 块映射与 I/O 提交的粒度失配", "pg": "18", "lvl": 1, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
            {"txt": "3.1.1 逐页处理带来的 CPU 开销", "pg": "18", "lvl": 2, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
            {"txt": "3.1.2 Large Folio 子区间多 BIO 提交与 I/O 生命周期失配", "pg": "19", "lvl": 2, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
            {"txt": "3.1.3 空洞文件与 I/O 区间划分的一致性", "pg": "22", "lvl": 2, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
        {"txt": "3.2 关键数据结构的冲突与兼容", "pg": "23", "lvl": 1, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
        {"txt": "3.2 f2fs_iomap_folio_state：folio->private 冲突与致命问题", "pg": "23", "lvl": 1, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
            {"txt": "3.2.1 功能背景", "pg": "23", "lvl": 2, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
            {"txt": "3.2.2 重大问题分析", "pg": "24", "lvl": 2, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
            {"txt": "3.2.2 I/O 描述符 (f2fs_io_info) 的粒度限制", "pg": "26", "lvl": 2, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
        {"txt": "3.3 压缩文件场景下的复杂性", "pg": "27", "lvl": 1, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
            {"txt": "3.3.1 压缩簇与 Folio 边界的不对齐", "pg": "27", "lvl": 2, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
            {"txt": "3.3.2 混合 I/O 模型（压缩与非压缩混排）", "pg": "28", "lvl": 2, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
        {"txt": "3.4 脏页回写中的写放大风险", "pg": "29", "lvl": 1, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
            {"txt": "3.4.1 F2FS 原有写回模型与写放大根源", "pg": "29", "lvl": 2, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
            {"txt": "3.4.2 Large Folio 与压缩写回的架构性挑战", "pg": "30", "lvl": 2, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},

        # 第四章
        {"txt": "第四章 基于 Folio 和 iomap 的方案设计与实现", "pg": "31", "lvl": 0, "fn": "黑体", "fs": S_4, "pg_fn": "宋体"},

        {"txt": "4.1 压缩文件 buffered read 的方案设计", "pg": "31", "lvl": 1, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
            {"txt": "4.1.1 设计目标与整体思路", "pg": "31", "lvl": 2, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
            {"txt": "4.1.2 核心迭代器：f2fs_compress_readpage_iter", "pg": "32", "lvl": 2, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
            {"txt": "4.1.3 统一的 Folio 完成路径：f2fs_iomap_finish_folio_read", "pg": "34", "lvl": 2, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
            {"txt": "4.1.4 解压完成回调与多 Folio 聚合：f2fs_decompress_end_io", "pg": "35", "lvl": 2, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
            {"txt": "4.1.5 流水线安全性与并发性分析", "pg": "37", "lvl": 2, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
        {"txt": "4.2 压缩文件 buffered write 的方案设计", "pg": "37", "lvl": 1, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
            {"txt": "4.2.1 算法整体思路：从“压缩簇”到“大 Folio”的视角切换", "pg": "37", "lvl": 2, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
            {"txt": "4.2.2 决赛方案核心代码：f2fs_buffered_write_iomap_begin", "pg": "38", "lvl": 2, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
            {"txt": "4.2.3 与 iomap_buffered_write 的协同与收益分析", "pg": "41", "lvl": 2, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
        {"txt": "4.3 普通/原子文件 buffered write 的统一实现", "pg": "42", "lvl": 1, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
            {"txt": "4.3.1 普通 buffered write：dirty 标记与空间预留的精细化", "pg": "42", "lvl": 2, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
            {"txt": "4.3.2 原子文件写入：f2fs_buffered_write_atomic_iomap_begin", "pg": "43", "lvl": 2, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
            {"txt": "4.3.3 写入失败与短写处理：从 f2fs_write_failed 到 iomap_write_delalloc_release", "pg": "45", "lvl": 2, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
        {"txt": "4.4 Page Writeback 子系统的改造与实现", "pg": "46", "lvl": 1, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
            {"txt": "4.4.1 设计目标与整体结构", "pg": "46", "lvl": 2, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
            {"txt": "4.4.2 f2fs_write_cache_folios：基于 writeback_iter 的高阶 Folio 写回", "pg": "47", "lvl": 2, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
            {"txt": "4.4.3 cc_bytes_pending：延迟解锁与生命周期耦合", "pg": "49", "lvl": 2, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
            {"txt": "4.4.4 f2fs_write_raw_pages：去掉危险解锁循环与重试集中化", "pg": "51", "lvl": 2, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
        {"txt": "4.5 垃圾回收（GC）路径的兼容与优化", "pg": "53", "lvl": 1, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
            {"txt": "4.5.1 问题回顾：GC 与 Large Folios / iomap 的“三重冲突”", "pg": "53", "lvl": 2, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
            {"txt": "4.5.2 f2fs_iomap_folio_state 的扩展：GC 状态与 cc_bytes_pending 的统一承载", "pg": "53", "lvl": 2, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
            {"txt": "4.5.3 move_data_page：GC 迁移路径中的精确脏标记", "pg": "55", "lvl": 2, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
            {"txt": "4.5.4 I/O 转向修复：f2fs_submit_page_bio 与 ra_data_block 的统一", "pg": "56", "lvl": 2, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},

        # 第五章
        {"txt": "第五章 项目性能测试", "pg": "57", "lvl": 0, "fn": "黑体", "fs": S_4, "pg_fn": "宋体"},

        {"txt": "5.1 测试环境与方法论", "pg": "57", "lvl": 1, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
            {"txt": "5.1.1 核心测试调度脚本：run_benchmark.py", "pg": "58", "lvl": 2, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
            {"txt": "5.1.2 底层测试执行脚本：fio_test.sh & clear_cache.sh", "pg": "61", "lvl": 2, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
        {"txt": "5.2 性能测试结果与分析", "pg": "63", "lvl": 1, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
            {"txt": "5.2.1 QEMU 虚拟机性能测试", "pg": "63", "lvl": 2, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
            {"txt": "5.2.2 树莓派 5 性能测试", "pg": "64", "lvl": 2, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
        {"txt": "5.3 数据处理与可视化脚本", "pg": "64", "lvl": 1, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
            {"txt": "5.3.1 日志解析脚本：parse_fio_logs.py", "pg": "64", "lvl": 2, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
            {"txt": "5.3.2 结果可视化脚本：visualize_results.py", "pg": "67", "lvl": 2, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},

        # 结尾
        {"txt": "结语", "pg": "55", "lvl": 0, "fn": "黑体", "fs": S_4, "pg_fn": "宋体"},
        {"txt": "参考文献", "pg": "56", "lvl": 0, "fn": "黑体", "fs": S_4, "pg_fn": "宋体"},
        {"txt": "致谢", "pg": "57", "lvl": 0, "fn": "黑体", "fs": S_4, "pg_fn": "宋体"},
    ]

    # --- 4. 循环生成 ---
    for item in toc_data:
        add_row(
            text=item["txt"],
            page_num=item["pg"],
            level=item["lvl"],
            font_name=item["fn"],
            font_size=item["fs"],
            page_font=item["pg_fn"],
            is_bold=False # 统一不加粗
        )

    # --- 5. 保存 ---
    doc.save('Final_Thesis_TOC.docx')
    print("生成完毕：Final_Thesis_TOC.docx")

if __name__ == "__main__":
    create_perfect_toc()


## 脚本使用提示词:
# 只允许更改toc_data,并且严格按照代码中的示例的格式
# 具体来说 如果是第一章这种大标题 必须{"txt": "第一章 引言", "pg": "1", "lvl": 0, "fn": "黑体", "fs": S_4, "pg_fn": "宋体"},
# 如果是1.1这种小标题 必须{"txt": "1.1 网站导航技术的国内外研究现状", "pg": "2", "lvl": 1, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
# 如果是1.1.1这种更小标题 必须{"txt": "1.1.1 JavaScript、CSS 和 HTC", "pg": "2", "lvl": 2, "fn": "宋体", "fs": S_X4, "pg_fn": "宋体"},
# 字典中的标题级别lvl都规定好了和标题级别数一一对应,并且字体也是按照示例中的严格规定好了
# 而且  # 摘要：黑体 四号，页码 I (宋体)
#        {"txt": "摘要", "pg": "I", "lvl": 0, "fn": "黑体", "fs": S_4, "pg_fn": "宋体"},

        # ABSTRACT：Times New Roman 四号，页码 II (宋体)
#        {"txt": "ABSTRACT", "pg": "II", "lvl": 0, "fn": "Times New Roman", "fs": S_4, "pg_fn": "宋体"},

# 和
 # 结尾：黑体 四号
        #{"txt": "结语", "pg": "104", "lvl": 0, "fn": "黑体", "fs": S_4, "pg_fn": "宋体"},
        #{"txt": "参考文献", "pg": "106", "lvl": 0, "fn": "黑体", "fs": S_4, "pg_fn": "宋体"},
        #{"txt": "致谢", "pg": "107", "lvl": 0, "fn": "黑体", "fs": S_4, "pg_fn": "宋体"},
# 必须严格保留并且是严格按照示例中的顺序一个在开头一个在结尾 我只要改动的是正文的目录结构。