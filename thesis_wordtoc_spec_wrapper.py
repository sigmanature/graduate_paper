#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate a Word-friendly thesis docx:
- Headings (level 1-3) are real Word headings (so TOC can be computed by Word)
- Inserts Chinese/English abstracts (per school appendix examples)
- Inserts a Word TOC field (levels 1-3, NO hyperlink to avoid blue)
- Leaves pagination to Microsoft Word (Windows) as the single source of truth
"""

from __future__ import annotations

from pathlib import Path

import thesis as T

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

BASE_DIR = Path(__file__).resolve().parent

# ---- You can edit these two titles if your school requires an exact one ----
CN_TITLE = "基于 Large Folio 与 iomap 的 F2FS 缓冲 I/O 优化研究"
EN_TITLE = "Optimizing F2FS Buffered I/O with Large Folios and iomap"

CN_ABSTRACT = '随着服务器内存容量与 NVMe 等高速存储设备的普及，Linux 内核长期以 4KB 页为基本管理单元的设计在页缓存与文件系统 I/O 路径上逐渐暴露出可扩展性瓶颈。为应对大内存与大块顺序 I/O 的需求，内核社区引入了 struct folio 及 Large Folios，并以 iomap 框架取代传统 buffer_head 模型，以更大的粒度管理页缓存并统一文件系统映射与写回逻辑。本文围绕 Large Folio 与 iomap 在页缓存路径中的落地展开，结合 F2FS 的日志结构、压缩簇与原子写等特性，分析大页缓存与细粒度块状态之间的粒度错配问题，并给出面向 F2FS 的工程化适配方案。在实现上，本文基于 iomap 的 buffered write / writeback 通路，完善 F2FS 对 large folios 的声明与分配，引入逐块脏页追踪与写回边界控制，确保在一个 Large Folio 内仅对实际变更的文件系统块提交 I/O，避免写放大；同时对失败回滚、压缩写回与原子文件路径中的关键钩子进行改造，使其与 iomap 的 DELALLOC 回收与状态机一致。实验结果表明，在顺序读写与典型工作负载下，该方案能够减少页缓存元数据与锁竞争开销，降低映射与回写路径上的 CPU 消耗，并在保持一致性的前提下提升 F2FS 的大块 I/O 吞吐与稳定性，为后续在 16KB/64KB 页配置与更大块设备场景中推广提供基础。'
CN_KEYWORDS = ["Large Folio", "Folio", "iomap", "F2FS", "页缓存", "写回"]

EN_ABSTRACT = 'With the widespread adoption of large-memory servers and high-performance NVMe storage, the long-standing 4KB page granularity in the Linux page cache and filesystem I/O stack increasingly limits scalability. To address large working sets and large sequential I/O, the kernel community introduced struct folio and Large Folios, and promoted the iomap framework to replace the legacy buffer_head-based model, enabling larger-granularity page-cache management and a unified mapping/writeback pipeline. This thesis investigates how Large Folios and iomap can be effectively integrated into the page-cache path, with a focus on F2FS and its log-structured design, compression clusters, and atomic write semantics. We analyze the granularity mismatch between large cache units and per-block filesystem state, and propose an engineering adaptation for F2FS. Based on iomap buffered write and writeback paths, our implementation enables F2FS to declare and allocate large folios, adds per-block dirty tracking and writeback boundary control so that only modified filesystem blocks within a large folio are written back, thereby mitigating write amplification. We further align error rollback, compressed writeback, and atomic file paths with iomap’s delalloc release and state-machine semantics. Experiments show that the proposed design reduces metadata and lock contention overhead in the page cache, lowers CPU cost on mapping and writeback paths, and improves throughput and stability for large-block I/O on F2FS while preserving consistency.'
EN_KEYWORDS = ["Large Folio", "folio", "iomap", "F2FS", "page cache", "writeback"]


def _set_run_font(run, eastasia: str, ascii_font: str, size_pt: float, *, bold=None, color_rgb=(0, 0, 0), underline=False):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), eastasia)
    run._element.rPr.rFonts.set(qn("w:ascii"), ascii_font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), ascii_font)
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    run.font.color.rgb = RGBColor(*color_rgb)
    run.font.underline = underline


def _apply_heading_styles(doc: Document):
    """Make Heading 1/2/3 visually identical to thesis.py, and force BLACK (avoid Word theme blue)."""
    # Heading 1: 三号宋体 16pt 居中 不加粗
    h1 = doc.styles["Heading 1"]
    h1.font.name = "宋体"
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    h1.font.size = Pt(16)
    h1.font.bold = False
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(0)
    h1.paragraph_format.line_spacing = 1.5

    # Heading 2: 四号宋体 14pt 加粗 居左
    h2 = doc.styles["Heading 2"]
    h2.font.name = "宋体"
    h2._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.space_before = Pt(0)
    h2.paragraph_format.space_after = Pt(0)
    h2.paragraph_format.line_spacing = 1.5

    # Heading 3: 小四宋体 12pt 居左
    h3 = doc.styles["Heading 3"]
    h3.font.name = "宋体"
    h3._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    h3.font.size = Pt(12)
    h3.font.bold = False
    h3.font.color.rgb = RGBColor(0, 0, 0)
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3.paragraph_format.space_before = Pt(0)
    h3.paragraph_format.space_after = Pt(0)
    h3.paragraph_format.line_spacing = 1.5


def _add_toc_field(doc: Document, levels: str = "1-3"):
    """Insert Word TOC field (NO \h hyperlink to avoid blue)."""
    p = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    # \z: hide page numbers in web layout; \u: use outline levels; \h intentionally omitted
    fld.set(qn("w:instr"), f'TOC \\o "{levels}" \\z \\u')
    p._p.append(fld)
    return p


def _center_line(doc: Document, text: str, *, font_ea="宋体", font_ascii="Times New Roman", size_pt=15, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    _set_run_font(r, font_ea, font_ascii, size_pt, bold=bold, color_rgb=(0, 0, 0))
    return p


def _body_line(doc: Document, text: str):
    # Abstract body uses same base as正文：小四宋体/Times，小四，首行缩进2字，1.5倍行距
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.first_line_indent = Pt(24)
    fmt.line_spacing = 1.5
    r = p.add_run(text)
    # reuse thesis.py's mixed-font rule
    try:
        T.set_cjk_west_font(r, Pt(12), bold=False)
    except Exception:
        _set_run_font(r, "宋体", "Times New Roman", 12, bold=False)
    return p


def _keywords_line_cn(doc: Document, kws: list[str]):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    r0 = p.add_run("关键词")
    _set_run_font(r0, "宋体", "Times New Roman", 12, bold=True)
    r1 = p.add_run("：" + "  ".join(kws))
    _set_run_font(r1, "宋体", "Times New Roman", 12, bold=False)
    return p


def _keywords_line_en(doc: Document, kws: list[str]):
    p = doc.add_paragraph()
    # In appendix sample, Key Words line uses hanging/negative indent; keep it simple & acceptable.
    p.paragraph_format.line_spacing = 1.5
    r0 = p.add_run("Key Words:")
    _set_run_font(r0, "Times New Roman", "Times New Roman", 12, bold=True)
    r1 = p.add_run(" " + "; ".join(kws))
    _set_run_font(r1, "Times New Roman", "Times New Roman", 12, bold=False)
    return p


def _patch_thesis_functions():
    """Monkey-patch thesis.py so headings become real Word headings (for TOC), and missing images won't crash."""
    orig_add_heading = T.add_heading
    orig_add_figure = T.add_figure

    def patched_add_heading(doc: Document, text: str, level: int):
        if level in (1, 2, 3):
            p = doc.add_paragraph(style=f"Heading {level}")
            r = p.add_run(text)
            # enforce black (avoid theme blue)
            if level == 1:
                _set_run_font(r, "宋体", "Times New Roman", 16, bold=False, color_rgb=(0, 0, 0))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif level == 2:
                _set_run_font(r, "宋体", "Times New Roman", 14, bold=True, color_rgb=(0, 0, 0))
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                _set_run_font(r, "宋体", "Times New Roman", 12, bold=False, color_rgb=(0, 0, 0))
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            return p
        return orig_add_heading(doc, text, level)

    def patched_add_figure(doc: Document, image_path: str, caption_text: str, chapter_no=None, width_cm=None, scale: float = 0.9):
        pth = Path(image_path)
        if not pth.is_absolute():
            cand = BASE_DIR / image_path
            if cand.exists():
                image_path = str(cand)

        if not Path(image_path).exists():
            # don't crash; keep pagination deterministic on your machine with real images
            try:
                T.add_placeholder(doc, f"图片缺失：{image_path}（{caption_text}）")
            except Exception:
                _body_line(doc, f"【图片缺失：{image_path}（{caption_text}）】")
            return None

        return orig_add_figure(doc, image_path, caption_text, chapter_no=chapter_no, width_cm=width_cm, scale=scale)

    T.add_heading = patched_add_heading
    T.add_figure = patched_add_figure


def build(out_docx: str = "Final_Thesis_WordTOC_Spec.docx"):
    _patch_thesis_functions()

    doc = Document()
    T.setup_page(doc)
    _apply_heading_styles(doc)

    # ============ Chinese Abstract ============
    _center_line(doc, CN_TITLE, font_ea="宋体", font_ascii="Times New Roman", size_pt=15, bold=False)
    _center_line(doc, "摘  要", font_ea="宋体", font_ascii="Times New Roman", size_pt=15, bold=False)
    _body_line(doc, CN_ABSTRACT)
    _keywords_line_cn(doc, CN_KEYWORDS)

    doc.add_page_break()

    # ============ English Abstract ============
    _center_line(doc, EN_TITLE, font_ea="Times New Roman", font_ascii="Times New Roman", size_pt=14, bold=False)
    _center_line(doc, "Abstract", font_ea="Times New Roman", font_ascii="Times New Roman", size_pt=14, bold=False)
    _body_line(doc, EN_ABSTRACT)
    _keywords_line_en(doc, EN_KEYWORDS)

    doc.add_page_break()

    # ============ TOC ============
    # Title: "目  录" (appendix sample uses 宋体三号)
    p_toc_title = doc.add_paragraph()
    p_toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_toc_title.add_run("目  录")
    _set_run_font(r, "宋体", "Times New Roman", 16, bold=False, color_rgb=(0, 0, 0))

    doc.add_paragraph()
    _add_toc_field(doc, levels="1-3")

    # New section: body starts at next page (page number restart is finalized in Word if needed)
    doc.add_section(WD_SECTION.NEW_PAGE)

    # ============ Main chapters ============
    T.generate_chapter_1(doc)
    T.generate_chapter_2(doc)
    T.generate_chapter_3(doc)
    T.generate_chapter_4(doc)
    T.generate_chapter_5(doc)

    # ============ Conclusion / References / Acknowledgements ============
    doc.add_page_break()
    T.add_heading(doc, "结语", level=1)
    T.add_body_paragraph(doc, "（此处填写结语内容，可在后续完善。）")

    doc.add_page_break()
    T.add_heading(doc, "参考文献", level=1)
    if hasattr(T, "REFERENCES") and hasattr(T, "add_reference_paragraph"):
        for ref in T.REFERENCES:
            T.add_reference_paragraph(doc, ref)

    doc.add_page_break()
    T.add_heading(doc, "致谢", level=1)
    T.add_body_paragraph(doc, "（此处填写致谢内容，可在后续完善。）")

    doc.save(out_docx)
    print(f"[OK] generated: {out_docx}")
    print("Next on Windows: run word_finalize_toc_export.py to update TOC, convert leader to '…' chars, and export PDF.")


if __name__ == "__main__":
    build()
