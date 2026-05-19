import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

import thesis as T

OUT_DOCX = "Final_Thesis_Stage1_WordTOC.docx"

CN_TITLE = "摘要"
EN_TITLE = "ABSTRACT"
TOC_TITLE = "目录"

CN_ABSTRACT = (
    "随着内存容量与高速存储设备性能持续提升，Linux 内核长期沿用的 4KB 基本内存管理单元逐渐暴露出扩展性瓶颈。"
    "小粒度页面会带来页表项数量庞大、TLB 覆盖率不足以及页缓存与 I/O 路径元数据开销偏高等问题，"
    "但本文并不对这一痛点展开铺陈，而是聚焦于大粒度页缓存管理在文件系统侧的工程落地。"
    "本文以 F2FS 为研究对象，围绕 Large Folio 与 iomap 两项关键机制展开分析与实现。"
    "其中，Large Folio 通过更大的页缓存管理粒度减少地址转换与页缓存管理开销，iomap 则通过区间式映射与每 Folio 状态管理"
    "为现代化 buffered I/O 路径提供统一抽象。针对 F2FS 在压缩文件、写回路径、GC 迁移以及 folio->private 状态复用等方面的难点，"
    "本文设计并实现了面向 Large Folio 的适配方案，重点解决了多 BIO 生命周期协调、压缩簇与 Folio 边界错位、逐块脏状态维护以及"
    " iomap 状态结构与 F2FS 私有标记的兼容问题。研究结果表明，该方案能够在保持文件系统语义正确性的前提下，"
    "提升大块顺序 I/O 场景下的路径效率，降低页缓存与回写阶段的额外开销，并为后续在 F2FS 中进一步推广更大页粒度与 iomap 化改造提供工程基础。"
)

EN_ABSTRACT = (
    "As memory capacity and storage bandwidth continue to grow, the long standing 4KB base memory management granularity in Linux "
    "has become increasingly restrictive for filesystem I/O paths. Instead of expanding on the general pain points of 4KB pages, "
    "this thesis focuses on the engineering integration of large grained page cache management in F2FS. The work centers on two key "
    "mechanisms, namely Large Folios and iomap. Large Folios improve cache management efficiency by using larger cache units, while iomap "
    "provides a modern interval based mapping model and per folio state tracking for buffered I/O. Targeting the difficulties in F2FS, "
    "including compressed file handling, page writeback, GC migration, and the conflict on folio private state usage, this thesis designs "
    "and implements a practical adaptation scheme for Large Folios. The implementation focuses on multi bio life cycle coordination, "
    "mismatched boundaries between compression clusters and folios, fine grained dirty state tracking, and the compatibility between iomap "
    "folio state and F2FS private markers. The results show that the proposed design can preserve correctness while improving large sequential "
    "I/O efficiency, reducing page cache and writeback overhead, and providing a solid foundation for future iomap based evolution in F2FS."
)

CN_KEYWORDS = ["Large Folio", "iomap", "F2FS", "页缓存", "文件系统"]
EN_KEYWORDS = ["Large Folio", "iomap", "F2FS", "page cache", "filesystem"]


PT_PER_CM = 72.0 / 2.54


def ensure_rfonts(run, east_asia="宋体", ascii_name="宋体"):
    run.font.name = ascii_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:ascii"), ascii_name)
    rfonts.set(qn("w:hAnsi"), ascii_name)


def set_para_outline_level(paragraph, level):
    p_pr = paragraph._p.get_or_add_pPr()
    outline = p_pr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        p_pr.append(outline)
    outline.set(qn("w:val"), str(level - 1))


def add_page_number_field(paragraph):
    run = paragraph.add_run()
    ensure_rfonts(run, "宋体", "宋体")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def set_section_page_number_format(section, fmt, start=None):
    sect_pr = section._sectPr
    pg = sect_pr.find(qn("w:pgNumType"))
    if pg is None:
        pg = OxmlElement("w:pgNumType")
        sect_pr.append(pg)
    pg.set(qn("w:fmt"), fmt)
    if start is not None:
        pg.set(qn("w:start"), str(start))


def add_footer_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    add_page_number_field(p)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    ensure_rfonts(run, "黑体", "黑体")
    run.font.size = Pt(16)
    run.bold = False
    set_para_outline_level(p, 1)
    return p


def add_song_paragraph(doc, text, first_line=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if first_line:
        pf.first_line_indent = Pt(24)
    run = p.add_run(text)
    ensure_rfonts(run, "宋体", "宋体")
    run.font.size = Pt(12)
    return p


def add_keywords(doc, label, words):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(6)
    pf.space_after = Pt(0)
    r1 = p.add_run(label)
    ensure_rfonts(r1, "黑体", "黑体")
    r1.font.size = Pt(12)
    r2 = p.add_run("；".join(words) if label.startswith("关键词") else ", ".join(words))
    ensure_rfonts(r2, "宋体", "宋体")
    r2.font.size = Pt(12)


def add_toc_field(paragraph):
    run = paragraph.add_run()
    ensure_rfonts(run, "宋体", "宋体")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \\o "1-3" \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = ""
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_text)
    run._r.append(fld_end)


def mark_existing_headings(doc):
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        if re.match(r"^第[一二三四五六七八九十]+章", text):
            set_para_outline_level(p, 1)
        elif re.match(r"^\d+\.\d+\s", text):
            set_para_outline_level(p, 2)
        elif re.match(r"^\d+\.\d+\.\d+\s", text):
            set_para_outline_level(p, 3)


def main():
    doc = Document()
    T.setup_page(doc)

    normal = doc.styles["Normal"]
    normal.font.size = Pt(12)
    normal.font.name = "宋体"
    try:
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        normal._element.rPr.rFonts.set(qn("w:ascii"), "宋体")
        normal._element.rPr.rFonts.set(qn("w:hAnsi"), "宋体")
    except Exception:
        pass

    # Front matter: Roman numerals
    set_section_page_number_format(doc.sections[0], "upperRoman", 1)
    add_footer_page_number(doc.sections[0])

    add_title(doc, CN_TITLE)
    add_song_paragraph(doc, CN_ABSTRACT, first_line=True)
    add_keywords(doc, "关键词：", CN_KEYWORDS)

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    set_section_page_number_format(doc.sections[-1], "upperRoman")
    add_footer_page_number(doc.sections[-1])

    add_title(doc, EN_TITLE)
    add_song_paragraph(doc, EN_ABSTRACT, first_line=True)
    add_keywords(doc, "Key Words: ", EN_KEYWORDS)

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    set_section_page_number_format(doc.sections[-1], "upperRoman")
    add_footer_page_number(doc.sections[-1])

    toc_title = add_title(doc, TOC_TITLE)
    toc_title.paragraph_format.space_after = Pt(3)

    begin_marker = doc.add_paragraph()
    rb = begin_marker.add_run("[[TOC_BEGIN]]")
    rb.font.hidden = True
    ensure_rfonts(rb, "宋体", "宋体")

    toc_p = doc.add_paragraph()
    toc_p.paragraph_format.line_spacing = 1.5
    add_toc_field(toc_p)

    end_marker = doc.add_paragraph()
    re_ = end_marker.add_run("[[TOC_END]]")
    re_.font.hidden = True
    ensure_rfonts(re_, "宋体", "宋体")

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    set_section_page_number_format(doc.sections[-1], "decimal", 1)
    add_footer_page_number(doc.sections[-1])

    for fn_name in [
        "generate_chapter_1",
        "generate_chapter_2",
        "generate_chapter_3",
        "generate_chapter_4",
        "generate_chapter_5",
    ]:
        fn = getattr(T, fn_name, None)
        if callable(fn):
            fn(doc)

    if hasattr(T, "REFERENCES") and hasattr(T, "add_reference_paragraph"):
        T.add_heading(doc, "参考文献", level=1)
        set_para_outline_level(doc.paragraphs[-1], 1)
        for idx, ref in enumerate(T.REFERENCES, start=1):
            T.add_reference_paragraph(doc, idx, ref)

    T.add_heading(doc, "致谢", level=1)
    set_para_outline_level(doc.paragraphs[-1], 1)
    add_song_paragraph(
        doc,
        "在本课题研究与论文撰写过程中，导师与课题组同学给予了大量指导与帮助。在此谨致谢意。",
        first_line=True,
    )

    mark_existing_headings(doc)
    out = Path(OUT_DOCX)
    doc.save(out)
    print(f"[OK] wrote {out}")


if __name__ == "__main__":
    main()
